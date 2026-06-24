import pandas as pd
import numpy as np
import xgboost as xgb
import shap
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import GroupShuffleSplit
from docx import Document
from docx.shared import Inches
import os

# Load data
df = pd.read_csv('real_greenflation_panel_data.csv')
cols = ['Inflation_Rate', 'Inflation_Lag1', 'CO2_Emissions_PC', 'GDP_Growth', 'Trade_Openness', 'Real_Interest_Rate']
df = df.dropna(subset=cols)

# 1. XGBoost & SHAP Analysis
X = df[['Inflation_Lag1', 'CO2_Emissions_PC', 'GDP_Growth', 'Trade_Openness', 'Real_Interest_Rate']]
y = df['Inflation_Rate']
groups = df['Country']

gss = GroupShuffleSplit(n_splits=1, train_size=0.8, random_state=42)
train_idx, test_idx = next(gss.split(X, y, groups))
X_train, X_test = X.iloc[train_idx], X.iloc[test_idx]
y_train, y_test = y.iloc[train_idx], y.iloc[test_idx]

model = xgb.XGBRegressor(n_estimators=200, max_depth=4, learning_rate=0.05, random_state=42)
model.fit(X_train, y_train)

# SHAP values
explainer = shap.TreeExplainer(model)
shap_values = explainer.shap_values(X_test)

# Plot 1: SHAP Summary Plot
plt.figure()
shap.summary_plot(shap_values, X_test, show=False)
plt.title('SHAP Summary Plot', pad=20)
plt.tight_layout()
plt.savefig('shap_summary.png', bbox_inches='tight')
plt.close()

# Plot 2: SHAP Dependence Plot for CO2
plt.figure()
shap.dependence_plot("CO2_Emissions_PC", shap_values, X_test, show=False, interaction_index=None)
plt.title('SHAP Dependence Plot: CO2 Emissions PC', pad=20)
plt.tight_layout()
plt.savefig('shap_dependence_co2.png', bbox_inches='tight')
plt.close()

# 2. Panel Data Graphs
# Plot 3: Time Trend (Macro trend)
df_time = df.groupby('Year')[['Inflation_Rate', 'CO2_Emissions_PC']].mean().reset_index()

fig, ax1 = plt.subplots(figsize=(10, 6))
color = 'tab:red'
ax1.set_xlabel('Year')
ax1.set_ylabel('Average Inflation Rate (%)', color=color)
ax1.plot(df_time['Year'], df_time['Inflation_Rate'], color=color, marker='o', label='Inflation Rate')
ax1.tick_params(axis='y', labelcolor=color)

ax2 = ax1.twinx()  
color = 'tab:blue'
ax2.set_ylabel('Average CO2 Emissions (Per Capita)', color=color)  
ax2.plot(df_time['Year'], df_time['CO2_Emissions_PC'], color=color, marker='s', linestyle='--', label='CO2 Emissions')
ax2.tick_params(axis='y', labelcolor=color)

plt.title('Macro Trend: Average Inflation vs Average CO2 Emissions (1990-2022)')
fig.tight_layout()  
plt.savefig('panel_time_trend.png', bbox_inches='tight')
plt.close()

# Plot 4: Panel Heterogeneity (Boxplot for top 15 countries by variance)
top_countries = df.groupby('Country')['Inflation_Rate'].var().nlargest(15).index
df_sub = df[df['Country'].isin(top_countries)]

plt.figure(figsize=(12, 6))
sns.boxplot(x='Country', y='Inflation_Rate', data=df_sub, palette='Set3', hue='Country', legend=False)
plt.xticks(rotation=45)
plt.title('Inflation Rate Heterogeneity (Selected 15 Countries)')
plt.ylabel('Inflation Rate (%)')
plt.xlabel('Country')
plt.tight_layout()
plt.savefig('panel_heterogeneity.png', bbox_inches='tight')
plt.close()

# 3. Update Word Document
doc_path = 'Greenflation_Findings.docx'
if os.path.exists(doc_path):
    doc = Document(doc_path)
    
    doc.add_heading('4. Advanced XGBoost & SHAP Diagnostics', level=1)
    doc.add_paragraph('To further dissect the non-linear relationship and the magnitude of the green transition shock, we employ SHAP (SHapley Additive exPlanations) values. The summary plot below demonstrates the directional impact of each feature, while the dependence plot isolates the effect of CO2 emissions on domestic inflation.')
    
    doc.add_heading('SHAP Summary Plot', level=2)
    doc.add_picture('shap_summary.png', width=Inches(6.0))
    
    doc.add_heading('SHAP Dependence Plot (CO2 Emissions)', level=2)
    doc.add_picture('shap_dependence_co2.png', width=Inches(6.0))
    
    doc.add_heading('5. Panel Data Visualization', level=1)
    doc.add_paragraph('The structural dynamics of the panel data are visualized through macro trends over time and cross-sectional heterogeneity.')
    
    doc.add_heading('Macro Trend: Global Average Inflation and CO2', level=2)
    doc.add_picture('panel_time_trend.png', width=Inches(6.0))
    
    doc.add_heading('Cross-Sectional Heterogeneity (Inflation Variance by Country)', level=2)
    doc.add_picture('panel_heterogeneity.png', width=Inches(6.0))
    
    doc.save(doc_path)
    print("Yeni grafikler olusturuldu ve Greenflation_Findings.docx dosyasina eklendi.")
else:
    print("Hata: Greenflation_Findings.docx bulunamadi.")
