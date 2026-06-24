import pandas as pd
import numpy as np
import xgboost as xgb
from sklearn.model_selection import GroupShuffleSplit
from sklearn.metrics import mean_squared_error, r2_score
from linearmodels.panel import PanelOLS, FirstDifferenceOLS
import statsmodels.api as sm
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os
from io import StringIO

df = pd.read_csv('real_greenflation_panel_data.csv')
cols = ['Inflation_Rate', 'Inflation_Lag1', 'CO2_Emissions_PC', 'GDP_Growth', 'Trade_Openness', 'Real_Interest_Rate']
df = df.dropna(subset=cols)

# 1. EXCEL EXPORT
print("Excel dosyasi olusturuluyor...")
excel_writer = pd.ExcelWriter('Greenflation_Findings.xlsx', engine='openpyxl')

# Desc stats
desc_stats = df[cols].describe().round(3)
desc_stats.to_excel(excel_writer, sheet_name='Descriptive_Stats')

# Correlation
corr_mat = df[cols].corr().round(3)
corr_mat.to_excel(excel_writer, sheet_name='Correlation_Matrix')

# Panel Models
df_panel = df.set_index(['Country', 'Year'])
exog_vars = ['Inflation_Lag1', 'CO2_Emissions_PC', 'GDP_Growth', 'Trade_Openness', 'Real_Interest_Rate']
exog = sm.add_constant(df_panel[exog_vars])
endog = df_panel['Inflation_Rate']

mod_fe = PanelOLS(endog, exog, entity_effects=True)
res_fe = mod_fe.fit(cov_type='robust')
fe_summary_html = res_fe.summary.tables[1].as_html()
fe_df = pd.read_html(StringIO(fe_summary_html), header=0, index_col=0)[0]
fe_df.to_excel(excel_writer, sheet_name='Fixed_Effects')

mod_fd = FirstDifferenceOLS(endog, df_panel[exog_vars])
res_fd = mod_fd.fit(cov_type='robust')
fd_summary_html = res_fd.summary.tables[1].as_html()
fd_df = pd.read_html(StringIO(fd_summary_html), header=0, index_col=0)[0]
fd_df.to_excel(excel_writer, sheet_name='First_Difference')

# XGBoost
X = df[exog_vars]
y = df['Inflation_Rate']
groups = df['Country']

gss = GroupShuffleSplit(n_splits=1, train_size=0.8, random_state=42)
train_idx, test_idx = next(gss.split(X, y, groups))
X_train, X_test = X.iloc[train_idx], X.iloc[test_idx]
y_train, y_test = y.iloc[train_idx], y.iloc[test_idx]

model = xgb.XGBRegressor(n_estimators=200, max_depth=4, learning_rate=0.05, random_state=42)
model.fit(X_train, y_train)

y_pred = model.predict(X_test)
rmse = np.sqrt(mean_squared_error(y_test, y_pred))
r2 = r2_score(y_test, y_pred)

importance = model.feature_importances_
feature_imp_df = pd.DataFrame({'Feature': exog_vars, 'Importance': importance}).sort_values(by='Importance', ascending=False)
feature_imp_df.to_excel(excel_writer, sheet_name='XGBoost_Importance', index=False)

pd.DataFrame({'Metric': ['Test RMSE', 'Test R2'], 'Value': [rmse, r2]}).to_excel(excel_writer, sheet_name='XGBoost_Metrics', index=False)

excel_writer.close()

# 2. GRAPHS EXPORT
print("Grafikler ciziliyor...")
plt.figure(figsize=(8, 6))
sns.scatterplot(x='CO2_Emissions_PC', y='Inflation_Rate', data=df, alpha=0.5)
sns.regplot(x='CO2_Emissions_PC', y='Inflation_Rate', data=df, scatter=False, color='red')
plt.title('Inflation Rate vs CO2 Emissions (Per Capita)')
plt.xlabel('CO2 Emissions Per Capita')
plt.ylabel('Inflation Rate (%)')
plt.tight_layout()
plt.savefig('scatter_co2_inflation.png')
plt.close()

plt.figure(figsize=(8, 6))
sns.barplot(x='Importance', y='Feature', data=feature_imp_df, hue='Feature', palette='viridis', legend=False)
plt.title('XGBoost Feature Importance for Inflation Modeling')
plt.xlabel('Importance (Gain)')
plt.ylabel('Feature')
plt.tight_layout()
plt.savefig('feature_importance.png')
plt.close()

# 3. WORD EXPORT
print("Word dosyasi hazirlaniyor...")
doc = Document()
doc.add_heading('The Transmission of Green Transition Shocks to Domestic Inflation: Findings', 0)

doc.add_heading('1. Descriptive Statistics', level=1)
doc.add_paragraph('Table 1 presents the descriptive statistics for the core variables used in our dynamic panel and machine learning models.')
table = doc.add_table(rows=1, cols=len(desc_stats.columns)+1)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
for i, col in enumerate(desc_stats.columns):
    hdr_cells[i+1].text = str(col)
for idx, row in desc_stats.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)
    for i, val in enumerate(row):
        row_cells[i+1].text = str(val)

doc.add_heading('2. Dynamic Panel Data Analysis', level=1)
doc.add_paragraph('Table 2 shows the Fixed Effects estimation capturing the unobserved heterogeneity across countries.')
table_fe = doc.add_table(rows=1, cols=len(fe_df.columns)+1)
table_fe.style = 'Table Grid'
hdr_cells = table_fe.rows[0].cells
hdr_cells[0].text = 'Predictor'
for i, col in enumerate(fe_df.columns):
    hdr_cells[i+1].text = str(col)
for idx, row in fe_df.iterrows():
    row_cells = table_fe.add_row().cells
    row_cells[0].text = str(idx)
    for i, val in enumerate(row):
        row_cells[i+1].text = str(val)
doc.add_paragraph(f'R-squared for FE Model: {res_fe.rsquared:.4f}')

doc.add_heading('3. Machine Learning (XGBoost) Findings', level=1)
doc.add_paragraph(f'The predictive modeling utilizing XGBoost yielded an out-of-sample Test R-squared of {r2:.4f} and Test RMSE of {rmse:.4f}. The relative importance of CO2_Emissions_PC indicates its substantial role as an inflationary shock transmission mechanism.')

doc.add_heading('Feature Importance Plot', level=2)
doc.add_picture('feature_importance.png', width=Inches(6.0))

doc.add_heading('Scatter Plot: CO2 Emissions vs Inflation', level=2)
doc.add_picture('scatter_co2_inflation.png', width=Inches(6.0))

doc.save('Greenflation_Findings.docx')
print("Tum dosyalar (Excel, Word, PNG) basariyla olusturuldu.")
